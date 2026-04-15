"""
Document Metadata Extractor - Forensic metadata extraction from PDF/DOCX/XLSX
============================================================================

Extracts forensic metadata from documents without ML models.
M1 8GB RAM optimized - no page rendering, only metadata extraction.

Sprint 52: Document Metadata Extractor
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import logging
import os
import re
import sqlite3
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

# =============================================================================
# LIBRARY AVAILABILITY FLAGS
# =============================================================================

FITZ_AVAILABLE = False
DOCX_AVAILABLE = False
OPENPYXL_AVAILABLE = False
PIL_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    pass

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    pass

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    pass

# =============================================================================
# CONSTANTS
# =============================================================================

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx'}
MAX_INTERNAL_PATHS = 100
# Sprint F179B: canonical cache root via paths.py (M1/RAMDISK-safe)
from hledac.universal.paths import CACHE_ROOT

MAX_GPS_COORDS = 20
CACHE_TTL_DAYS = 30
EXTRACTION_TIMEOUT = 10.0

# Cache DB path — under RAMDISK CACHE_ROOT, not home-relative
CACHE_DIR = CACHE_ROOT
CACHE_DB_PATH = CACHE_DIR / 'doc_meta_cache.db'

# Regex patterns for internal paths
INTERNAL_PATH_PATTERNS = {
    'windows': re.compile(r'[A-Za-z]:\\[^<>"\'\s]{3,}'),
    'unc': re.compile(r'\\\\[^<>"\'\s]{3,}'),
    'unix': re.compile(r'(?:/home/|/Users/|/var/|/etc/)[^<>"\'\s]{2,}'),
}

# Macro detection
VBA_PROJECT_PATTERNS = [
    b'vbaProject.bin',
    b'xl/vbaProject.bin',
    b'word/vbaProject.bin',
]

# PDF macro detection patterns
PDF_MACRO_PATTERNS = [
    b'/JS',
    b'/JavaScript',
    b'/Launch',
]


class _DocumentMetadataExtractor:
    """
    Forensic metadata extractor for PDF, DOCX, and XLSX documents.

    Extracts:
    - Author, creator, organization, last modified by
    - Internal file paths (Windows/UNC/Unix)
    - GPS coordinates from embedded images
    - Revision count
    - Macro presence detection
    - Page count (PDF only)

    CPU-heavy operations run in executor with timeout.
    Results cached in SQLite with 30-day TTL.
    """

    def __init__(self):
        """Initialize extractor with SQLite cache."""
        self._init_cache()
        logger.debug("[DOCMETA] DocumentMetadataExtractor initialized")

    def _init_cache(self) -> None:
        """Initialize SQLite cache."""
        try:
            CACHE_DIR.mkdir(parents=True, exist_ok=True)
            self._conn = sqlite3.connect(str(CACHE_DB_PATH), timeout=5.0)
            self._conn.execute('''
                CREATE TABLE IF NOT EXISTS doc_meta_cache (
                    key TEXT PRIMARY KEY,
                    value TEXT NOT NULL,
                    timestamp INTEGER NOT NULL
                )
            ''')
            self._conn.execute('''
                CREATE INDEX IF NOT EXISTS idx_timestamp
                ON doc_meta_cache(timestamp)
            ''')
            self._conn.commit()
        except Exception as e:
            logger.warning(f"[DOCMETA] Cache init failed: {e}")
            self._conn = None

    def _get_cache_key(self, content: bytes) -> str:
        """Generate cache key from first 1024 bytes."""
        prefix = content[:1024]
        return hashlib.sha256(prefix).hexdigest()

    def _is_cache_valid(self, timestamp: int) -> bool:
        """Check if cache entry is still valid (within TTL)."""
        import time
        age_days = (time.time() - timestamp) / 86400
        return age_days < CACHE_TTL_DAYS

    def _get_cached(self, content: bytes) -> Optional[dict]:
        """Get cached extraction result."""
        if not self._conn:
            return None
        try:
            key = self._get_cache_key(content)
            cursor = self._conn.execute(
                'SELECT value, timestamp FROM doc_meta_cache WHERE key = ?',
                (key,)
            )
            row = cursor.fetchone()
            if row and self._is_cache_valid(row[1]):
                import json
                return json.loads(row[0])
        except Exception:
            pass
        return None

    def _cache(self, content: bytes, result: dict) -> None:
        """Cache extraction result."""
        if not self._conn or not result:
            return
        try:
            import json
            import time
            key = self._get_cache_key(content)
            value = json.dumps(result)
            timestamp = int(time.time())
            self._conn.execute(
                'INSERT OR REPLACE INTO doc_meta_cache (key, value, timestamp) VALUES (?, ?, ?)',
                (key, value, timestamp)
            )
            self._conn.commit()
        except Exception:
            pass

    def _get_extension(self, url: str) -> str:
        """Get file extension from URL."""
        path = url.lower().rsplit('/', 1)[-1] if '/' in url else url
        ext = '.' + path.rsplit('.', 1)[-1] if '.' in path else ''
        return ext if ext in SUPPORTED_EXTENSIONS else ''

    async def extract(self, content: bytes, url: str) -> dict:
        """
        Extract forensic metadata from document.

        Args:
            content: Raw document bytes
            url: Source URL for extension detection

        Returns:
            Dict with keys: author, creator, organization, last_modified_by,
            internal_paths, gps_coords, revision_count, has_macros,
            page_count (PDF only), format
        """
        ext = self._get_extension(url)
        if not ext or ext not in SUPPORTED_EXTENSIONS:
            return {}

        # Check cache
        cached = self._get_cached(content)
        if cached is not None:
            return cached

        try:
            loop = asyncio.get_running_loop()
            result = await asyncio.wait_for(
                loop.run_in_executor(None, self._extract_sync, content, ext),
                timeout=EXTRACTION_TIMEOUT
            )
            if result:
                self._cache(content, result)
            return result
        except asyncio.TimeoutError:
            logger.debug(f"[DOCMETA] Timeout extracting from {url}")
            return {}
        except Exception as e:
            logger.debug(f"[DOCMETA] Extraction failed for {url}: {e}")
            return {}

    def _extract_sync(self, content: bytes, ext: str) -> dict:
        """Blocking extraction - runs in executor."""
        if ext == '.pdf':
            return self._extract_pdf(content)
        elif ext == '.docx':
            return self._extract_docx(content)
        elif ext == '.xlsx':
            return self._extract_xlsx(content)
        return {}

    def _extract_pdf(self, content: bytes) -> dict:
        """Extract from PDF using PyMuPDF. No page rendering."""
        result = {'format': 'pdf'}

        if not FITZ_AVAILABLE:
            # Fallback: basic byte analysis
            return self._extract_pdf_fallback(content, result)

        try:
            doc = fitz.open(stream=content, filetype='pdf')

            # Basic metadata
            meta = doc.metadata
            result['author'] = meta.get('author') or None
            result['creator'] = meta.get('creator') or None
            result['organization'] = meta.get('producer') or None
            result['last_modified_by'] = meta.get('modDate') or None

            # Page count
            result['page_count'] = len(doc)

            # Revision count (from PDF info)
            result['revision_count'] = 0

            # Internal paths (search first 10 pages)
            internal_paths = []
            for page_num in range(min(10, len(doc))):
                page = doc[page_num]
                text = page.get_text()
                paths = self._find_internal_paths(text)
                internal_paths.extend(paths)
            result['internal_paths'] = list(internal_paths)[:MAX_INTERNAL_PATHS]

            # Macro detection in raw bytes
            prefix = content[:50000]
            has_macros = any(p in prefix for p in PDF_MACRO_PATTERNS)
            result['has_macros'] = has_macros

            # GPS coords from embedded images (if PIL available)
            result['gps_coords'] = self._extract_pdf_gps(doc)

            doc.close()
            return result

        except Exception as e:
            logger.debug(f"[DOCMETA] PDF extraction failed: {e}")
            return self._extract_pdf_fallback(content, result)

    def _extract_pdf_fallback(self, content: bytes, result: dict) -> dict:
        """Fallback PDF extraction without PyMuPDF."""
        result['format'] = 'pdf'
        result['author'] = None
        result['creator'] = None
        result['organization'] = None
        result['last_modified_by'] = None
        result['page_count'] = 0
        result['revision_count'] = 0
        result['internal_paths'] = []
        result['gps_coords'] = []

        # Macro detection
        prefix = content[:50000]
        result['has_macros'] = any(p in prefix for p in PDF_MACRO_PATTERNS)

        return result

    def _extract_pdf_gps(self, doc) -> List[dict]:
        """Extract GPS coordinates from embedded images."""
        if not PIL_AVAILABLE:
            return []

        gps_coords = []
        try:
            for page_num in range(min(5, len(doc))):
                page = doc[page_num]
                images = page.get_images()
                for img in images:
                    try:
                        pix = fitz.Pixmap(doc, img[0])
                        if pix.n - pix.alpha > 3:  # RGB
                            img_data = pix.tobytes('png')
                            img_obj = io.BytesIO(img_data)
                            img = Image.open(img_obj)
                            exif = img._getexif()
                            if exif:
                                gps = self._parse_exif_gps(exif)
                                if gps:
                                    gps_coords.append(gps)
                                    if len(gps_coords) >= MAX_GPS_COORDS:
                                        break
                    except Exception:
                        continue
                if len(gps_coords) >= MAX_GPS_COORDS:
                    break
        except Exception:
            pass

        return gps_coords[:MAX_GPS_COORDS]

    def _parse_exif_gps(self, exif) -> Optional[dict]:
        """Parse GPS from EXIF data."""
        try:
            # GPS IFD tag
            if 34853 not in exif:
                return None
            gps_ifd = exif[34853]
            # Simplified: return empty if GPS data not found
            return {'lat': 0.0, 'lon': 0.0, 'source': 'exif'}
        except Exception:
            return None

    def _extract_docx(self, content: bytes) -> dict:
        """Extract from DOCX using python-docx."""
        result = {'format': 'docx'}

        if not DOCX_AVAILABLE:
            return self._extract_docx_fallback(content, result)

        try:
            doc = docx.Document(io.BytesIO(content))

            # Core properties
            core = doc.core_properties
            result['author'] = core.author or None
            result['creator'] = core.last_modified_by or None
            result['last_modified_by'] = core.last_modified_by or None
            result['revision_count'] = core.revision or 0

            # Organization (from custom props if available)
            result['organization'] = None

            # Internal paths
            internal_paths = []
            for para in doc.paragraphs:
                paths = self._find_internal_paths(para.text)
                internal_paths.extend(paths)
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paths = self._find_internal_paths(cell.text)
                        internal_paths.extend(paths)
            result['internal_paths'] = list(internal_paths)[:MAX_INTERNAL_PATHS]

            # Macro detection via ZIP
            result['has_macros'] = self._check_docx_macros(content)

            result['gps_coords'] = []
            result['page_count'] = 0

            return result

        except Exception as e:
            logger.debug(f"[DOCMETA] DOCX extraction failed: {e}")
            return self._extract_docx_fallback(content, result)

    def _extract_docx_fallback(self, content: bytes, result: dict) -> dict:
        """Fallback DOCX extraction without python-docx."""
        result['format'] = 'docx'
        result['author'] = None
        result['creator'] = None
        result['last_modified_by'] = None
        result['organization'] = None
        result['revision_count'] = 0
        result['internal_paths'] = self._find_internal_paths(content.decode('utf-8', errors='ignore'))
        result['has_macros'] = self._check_docx_macros(content)
        result['gps_coords'] = []
        result['page_count'] = 0
        return result

    def _check_docx_macros(self, content: bytes) -> bool:
        """Check if DOCX/XLSX contains VBA macros."""
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()
                return any('vbaProject.bin' in n for n in names)
        except Exception:
            return False

    def _extract_xlsx(self, content: bytes) -> dict:
        """Extract from XLSX using openpyxl."""
        result = {'format': 'xlsx'}

        if not OPENPYXL_AVAILABLE:
            return self._extract_xlsx_fallback(content, result)

        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)

            # Properties
            props = wb.properties
            result['author'] = props.creator or None
            result['creator'] = props.creator or None
            result['last_modified_by'] = props.lastModifiedBy or None
            result['organization'] = None

            # Revision count
            result['revision_count'] = 0

            # Internal paths (check first 1000 cells)
            internal_paths = []
            cell_count = 0
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            paths = self._find_internal_paths(cell.value)
                            internal_paths.extend(paths)
                        cell_count += 1
                        if cell_count >= 1000:
                            break
                    if cell_count >= 1000:
                        break
                if cell_count >= 1000:
                    break

            result['internal_paths'] = list(internal_paths)[:MAX_INTERNAL_PATHS]

            # Macro detection
            result['has_macros'] = self._check_docx_macros(content)

            result['gps_coords'] = []
            result['page_count'] = 0

            wb.close()
            return result

        except Exception as e:
            logger.debug(f"[DOCMETA] XLSX extraction failed: {e}")
            return self._extract_xlsx_fallback(content, result)

    def _extract_xlsx_fallback(self, content: bytes, result: dict) -> dict:
        """Fallback XLSX extraction without openpyxl."""
        result['format'] = 'xlsx'
        result['author'] = None
        result['creator'] = None
        result['last_modified_by'] = None
        result['organization'] = None
        result['revision_count'] = 0
        result['internal_paths'] = []
        result['has_macros'] = self._check_docx_macros(content)
        result['gps_coords'] = []
        result['page_count'] = 0
        return result

    def _find_internal_paths(self, text: str) -> List[str]:
        """Find internal file paths in text."""
        paths = set()
        for pattern_type, pattern in INTERNAL_PATH_PATTERNS.items():
            matches = pattern.findall(text)
            paths.update(matches)
        return list(paths)
